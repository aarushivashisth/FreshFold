"""
Generate a comprehensive Word document explaining the FreshFold Laundry Management System.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ─── Helper Functions ─────────────────────────────────────────────────────────

def set_run_font(run, size=11, bold=False, color=None, font_name="Calibri"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(text, size=26):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=(60, 38, 32))

def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(90, 58, 50)

def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p

def add_code_block(title, code, language=""):
    add_para(f"📄 {title}", bold=True, size=11)
    p = doc.add_paragraph()
    run = p.add_run(code)
    set_run_font(run, size=9, font_name="Consolas", color=(30, 30, 30))
    p.paragraph_format.left_indent = Inches(0.3)

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run2 = p.add_run(text)
        set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()  # spacer
doc.add_paragraph()
doc.add_paragraph()

add_title("FreshFold", size=36)
add_title("Laundry Management System", size=22)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Complete Project Documentation")
set_run_font(run, size=16, color=(90, 58, 50))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A full-stack web application for online laundry service management\nbuilt with HTML, CSS, JavaScript (Frontend) and Node.js, Express, MySQL (Backend)")
set_run_font(run, size=12, color=(100, 100, 100))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("Table of Contents", level=1)

toc_items = [
    "1. Project Overview",
    "2. Technology Stack",
    "3. Project Directory Structure",
    "4. Database Design (db.sql)",
    "5. Backend – Server Configuration",
    "   5.1 package.json",
    "   5.2 db.js – Database Connection",
    "   5.3 server.js – Express Server Entry Point",
    "   5.4 seedAdmin.js – Admin Seeder Script",
    "6. Backend – API Routes",
    "   6.1 userRoutes.js – User Authentication",
    "   6.2 orderRoutes.js – Order Management",
    "   6.3 feedbackRoutes.js – Feedback System",
    "7. Frontend Pages",
    "   7.1 index.html – Entry Redirect",
    "   7.2 login.html – User Login",
    "   7.3 register.html – User Registration",
    "   7.4 home.html – Home / Landing Page",
    "   7.5 services.html – Service Selection & Cart",
    "   7.6 receipt.html – Order Receipt & Checkout",
    "   7.7 payment.html – Payment Processing",
    "   7.8 order-status.html – Order Tracking",
    "   7.9 feedback.html – Customer Feedback",
    "   7.10 admin.html – Admin Dashboard",
    "8. Application Flow & User Journey",
    "9. API Endpoints Reference",
    "10. Environment & Configuration",
    "11. How to Run the Project",
    "12. Key Features Summary",
]

for item in toc_items:
    add_para(item, size=11)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("1. Project Overview", level=1)

add_para(
    "FreshFold is a full-stack web-based Laundry Management System designed to allow customers "
    "to browse laundry services, select clothing types and quantities, place orders, track order status, "
    "make payments, and submit feedback. The system also includes an Admin Dashboard where the business "
    "owner can manage all orders and view customer feedback."
)

add_para(
    "The project follows a client-server architecture with a clear separation between the frontend "
    "(static HTML/CSS/JS pages) and the backend (Node.js/Express REST API). Data is persisted in a "
    "MySQL relational database."
)

add_heading_custom("Key Objectives", level=2)
add_bullet("Allow customers to register, login, and manage their laundry orders online.")
add_bullet("Provide an interactive service selection interface with quantity controls and service checklists.")
add_bullet("Generate itemized receipts with GST and delivery charge calculations.")
add_bullet("Support multiple payment modes: UPI, Credit/Debit Card, and Cash on Delivery (COD).")
add_bullet("Enable real-time order tracking with a visual step-by-step status indicator.")
add_bullet("Collect customer feedback with a star-rating system.")
add_bullet("Provide an Admin Dashboard to manage orders and view feedback.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("2. Technology Stack", level=1)

add_heading_custom("Frontend", level=2)
add_bullet("HTML5 – Semantic page structure")
add_bullet("CSS3 – Inline styles with glassmorphism, animations, and responsive design")
add_bullet("Vanilla JavaScript – DOM manipulation, Fetch API for AJAX calls, localStorage for session management")
add_bullet("Font Awesome – Icon library for social media icons")
add_bullet("Google Fonts (Playfair Display) – Premium typography")

add_heading_custom("Backend", level=2)
add_bullet("Node.js – JavaScript runtime environment")
add_bullet("Express.js – Web framework for building REST APIs")
add_bullet("MySQL2 – MySQL driver with Promise support")
add_bullet("bcrypt / bcryptjs – Password hashing library")
add_bullet("cors – Cross-Origin Resource Sharing middleware")
add_bullet("dotenv – Environment variable management")
add_bullet("jsonwebtoken – JWT library (installed but not actively used in current version)")

add_heading_custom("Database", level=2)
add_bullet("MySQL – Relational database for persistent data storage")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. DIRECTORY STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("3. Project Directory Structure", level=1)

structure = """FreshFold/
├── .gitignore                  # Git ignore rules
├── README.md                   # Project readme
├── db.sql                      # Database schema & seed data
│
├── backend/
│   ├── .env                    # Environment variables (DB credentials)
│   ├── package.json            # Node.js dependencies & scripts
│   ├── package-lock.json       # Locked dependency versions
│   ├── db.js                   # MySQL connection pool
│   ├── server.js               # Express server entry point
│   ├── seedAdmin.js            # Script to create admin user
│   ├── node_modules/           # Installed packages
│   └── routes/
│       ├── userRoutes.js       # /api/users – Register, Login, Get User
│       ├── orderRoutes.js      # /api/orders – Create, List, Update Status
│       └── feedbackRoutes.js   # /api/feedback – Create, List All
│
└── frontend/
    ├── index.html              # Redirect to login page
    ├── login.html              # User login page
    ├── register.html           # User registration page
    ├── home.html               # Home / landing page
    ├── services.html           # Service selection & cart
    ├── receipt.html            # Order receipt with billing
    ├── payment.html            # Payment method selection
    ├── order-status.html       # Order tracking page
    ├── feedback.html           # Customer feedback form
    ├── admin.html              # Admin dashboard
    ├── 1.png                   # Logo image
    ├── Black and White Initials elegant cursive logo.png
    └── Cream and Brown Photographic Beauty Site Launch Website.png  # Background image
"""

add_code_block("Directory Tree", structure)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. DATABASE DESIGN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("4. Database Design (db.sql)", level=1)

add_para(
    "The database is named 'freshfold' and contains 6 tables that store all application data. "
    "The schema uses foreign key relationships with ON DELETE CASCADE for data integrity."
)

add_heading_custom("4.1 users Table", level=2)
add_para("Stores registered user accounts including both customers and admins.")

# Table definition
table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
headers = ["Column", "Type", "Description"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

user_cols = [
    ("id", "INT AUTO_INCREMENT PK", "Unique user identifier"),
    ("name", "VARCHAR(200)", "Full name of the user"),
    ("email", "VARCHAR(200) UNIQUE", "User email address"),
    ("username", "VARCHAR(100) UNIQUE", "Login username"),
    ("password", "VARCHAR(255)", "Bcrypt hashed password"),
    ("address", "TEXT", "Delivery address"),
    ("phone", "VARCHAR(30)", "Phone number"),
    ("role", "ENUM('customer','admin')", "User role, defaults to 'customer'"),
]
for i, (col, typ, desc) in enumerate(user_cols):
    table.rows[i+1].cells[0].text = col
    table.rows[i+1].cells[1].text = typ
    table.rows[i+1].cells[2].text = desc

doc.add_paragraph()  # spacer

add_heading_custom("4.2 services Table", level=2)
add_para("Catalog of available laundry services with pricing.")
table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

svc_cols = [
    ("id", "INT AUTO_INCREMENT PK", "Unique service identifier"),
    ("name", "VARCHAR(100)", "Service name (e.g. Washing, Ironing)"),
    ("description", "TEXT", "Service description"),
    ("price_per_unit", "DECIMAL(10,2)", "Price per unit"),
]
for i, (col, typ, desc) in enumerate(svc_cols):
    table2.rows[i+1].cells[0].text = col
    table2.rows[i+1].cells[1].text = typ
    table2.rows[i+1].cells[2].text = desc

doc.add_paragraph()

add_heading_custom("4.3 orders Table", level=2)
add_para("Stores customer orders with payment and status tracking.")
table3 = doc.add_table(rows=10, cols=3)
table3.style = 'Table Grid'
for i, h in enumerate(headers):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

order_cols = [
    ("id", "INT AUTO_INCREMENT PK", "Unique order identifier"),
    ("user_id", "INT FK → users(id)", "Customer who placed the order"),
    ("total_amount", "DECIMAL(10,2)", "Total order amount"),
    ("payment_mode", "VARCHAR(30)", "Payment method (COD, UPI, Card)"),
    ("payment_status", "VARCHAR(50)", "Payment status (pending/collected)"),
    ("status", "VARCHAR(50)", "Order status (placed, picked, etc.)"),
    ("pickup_slot", "VARCHAR(100)", "Scheduled pickup time"),
    ("current_step", "INT DEFAULT 0", "Tracking progress step (0-6)"),
    ("created_at / updated_at", "TIMESTAMP", "Order creation and last update time"),
]
for i, (col, typ, desc) in enumerate(order_cols):
    table3.rows[i+1].cells[0].text = col
    table3.rows[i+1].cells[1].text = typ
    table3.rows[i+1].cells[2].text = desc

doc.add_paragraph()

add_heading_custom("4.4 order_items Table", level=2)
add_para("Stores individual line items within an order (clothing type, quantity, selected services, rate).")
table4 = doc.add_table(rows=6, cols=3)
table4.style = 'Table Grid'
for i, h in enumerate(headers):
    table4.rows[0].cells[i].text = h
    for p in table4.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

oi_cols = [
    ("id", "INT AUTO_INCREMENT PK", "Unique item identifier"),
    ("order_id", "INT FK → orders(id)", "Parent order reference"),
    ("item_type", "VARCHAR(100)", "Clothing type (Cotton, Silk, etc.)"),
    ("quantity", "INT", "Number of items"),
    ("services", "VARCHAR(255)", "Comma-separated service names"),
]
for i, (col, typ, desc) in enumerate(oi_cols):
    table4.rows[i+1].cells[0].text = col
    table4.rows[i+1].cells[1].text = typ
    table4.rows[i+1].cells[2].text = desc

doc.add_paragraph()

add_heading_custom("4.5 feedback Table", level=2)
add_para("Stores customer feedback and ratings.")
table5 = doc.add_table(rows=6, cols=3)
table5.style = 'Table Grid'
for i, h in enumerate(headers):
    table5.rows[0].cells[i].text = h
    for p in table5.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

fb_cols = [
    ("id", "INT AUTO_INCREMENT PK", "Unique feedback identifier"),
    ("user_id", "INT FK → users(id)", "User who submitted feedback"),
    ("order_id", "INT NULL", "Related order (optional)"),
    ("message", "TEXT", "Feedback message text"),
    ("rating", "INT NULL", "Star rating (1-5)"),
]
for i, (col, typ, desc) in enumerate(fb_cols):
    table5.rows[i+1].cells[0].text = col
    table5.rows[i+1].cells[1].text = typ
    table5.rows[i+1].cells[2].text = desc

doc.add_paragraph()

add_heading_custom("4.6 Seed Data", level=2)
add_para("The db.sql file also includes INSERT statements to populate the services table with 4 default services:")
add_bullet("Washing – ₹40 per unit")
add_bullet("Ironing – ₹20 per unit")
add_bullet("Dry Clean – ₹80 per unit")
add_bullet("Premium Wash – ₹60 per unit")

add_para("Full db.sql Code:", bold=True)
db_sql = open(os.path.join(os.path.dirname(__file__), "db.sql"), "r", encoding="utf-8").read()
add_code_block("db.sql", db_sql, "sql")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. BACKEND – SERVER CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("5. Backend – Server Configuration", level=1)

# 5.1 package.json
add_heading_custom("5.1 package.json", level=2)
add_para(
    "This file defines the Node.js project metadata, entry point, scripts, and dependencies. "
    "It is located in the backend/ directory."
)

pkg_json = open(os.path.join(os.path.dirname(__file__), "backend", "package.json"), "r", encoding="utf-8").read()
add_code_block("backend/package.json", pkg_json, "json")

add_para("Key Dependencies Explained:", bold=True)
add_bullet("express", bold_prefix="express: ")
add_bullet(" – The web framework that handles HTTP routing and middleware.")
add_bullet("mysql2", bold_prefix="mysql2: ")
add_bullet(" – MySQL driver with native Promise/async-await support.")
add_bullet("bcrypt / bcryptjs", bold_prefix="bcrypt: ")
add_bullet(" – Libraries to securely hash and compare passwords.")
add_bullet("cors", bold_prefix="cors: ")
add_bullet(" – Middleware to allow cross-origin requests from the frontend.")
add_bullet("dotenv", bold_prefix="dotenv: ")
add_bullet(" – Loads environment variables from a .env file.")
add_bullet("jsonwebtoken", bold_prefix="jsonwebtoken: ")
add_bullet(" – JWT library (available for future authentication enhancement).")

add_para("Scripts:", bold=True)
add_bullet('"npm start" runs "node server.js" to start the server.')
add_bullet('"npm run seed" runs "node seedAdmin.js" to create the admin user.')

doc.add_paragraph()

# 5.2 db.js
add_heading_custom("5.2 db.js – Database Connection", level=2)
add_para(
    "This file creates a MySQL connection pool using the mysql2 library. It reads database credentials "
    "from environment variables (via dotenv) with fallback defaults. The pool is exported as a Promise-based "
    "interface for async/await usage throughout the application."
)

add_para("How it works:", bold=True)
add_bullet("Creates a connection pool with a maximum of 10 simultaneous connections.")
add_bullet("Uses environment variables DB_HOST, DB_USER, DB_PASSWORD, DB_NAME for configuration.")
add_bullet("Falls back to localhost/root/1707/freshfold if .env is not configured.")
add_bullet("Exports pool.promise() so all queries can use async/await.")

db_js = open(os.path.join(os.path.dirname(__file__), "backend", "db.js"), "r", encoding="utf-8").read()
add_code_block("backend/db.js", db_js, "javascript")

doc.add_paragraph()

# 5.3 server.js
add_heading_custom("5.3 server.js – Express Server Entry Point", level=2)
add_para(
    "This is the main entry point of the backend application. It sets up the Express server, "
    "connects middleware, registers API routes, serves the frontend static files, and starts "
    "listening on port 5000."
)

add_para("Functionality breakdown:", bold=True)
add_bullet("Imports and configures Express, CORS, dotenv, and the database pool.")
add_bullet("Tests the MySQL connection on startup – exits with an error if the DB is unreachable.")
add_bullet("Serves the frontend/ directory as static files (HTML, CSS, JS, images).")
add_bullet("Mounts three route groups: /api/users, /api/orders, /api/feedback.")
add_bullet("The root route (/) redirects to login.html as the default landing page.")
add_bullet("Listens on PORT from environment variables, defaulting to 5000.")

server_js = open(os.path.join(os.path.dirname(__file__), "backend", "server.js"), "r", encoding="utf-8").read()
add_code_block("backend/server.js", server_js, "javascript")

doc.add_paragraph()

# 5.4 seedAdmin.js
add_heading_custom("5.4 seedAdmin.js – Admin Seeder Script", level=2)
add_para(
    "This is a one-time utility script to create the initial admin user in the database. "
    "It is run via 'npm run seed' before the application is first used."
)

add_para("What it does:", bold=True)
add_bullet("Hashes the default admin password ('12345') using bcrypt with a salt round of 8.")
add_bullet("Checks if an admin user already exists (by username 'admin' or email 'admin@freshfold.com').")
add_bullet("If no admin exists, inserts a new user with role='admin'.")
add_bullet("The created admin credentials are: username='admin1', password='12345'.")

seed_js = open(os.path.join(os.path.dirname(__file__), "backend", "seedAdmin.js"), "r", encoding="utf-8").read()
add_code_block("backend/seedAdmin.js", seed_js, "javascript")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. BACKEND – API ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("6. Backend – API Routes", level=1)

add_para(
    "The backend exposes RESTful API endpoints organized into three route modules. "
    "Each module is a function that receives the database pool and returns an Express Router."
)

# 6.1 userRoutes.js
add_heading_custom("6.1 userRoutes.js – User Authentication", level=2)
add_para("File: backend/routes/userRoutes.js", bold=True)
add_para("This module handles user registration, login, and profile retrieval.")

add_para("Endpoints:", bold=True)

t = doc.add_table(rows=4, cols=4)
t.style = 'Table Grid'
t_headers = ["Method", "Endpoint", "Description", "Auth Required"]
for i, h in enumerate(t_headers):
    t.rows[0].cells[i].text = h
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

routes_data = [
    ("POST", "/api/users/register", "Register a new customer account", "No"),
    ("POST", "/api/users/login", "Login with username & password", "No"),
    ("GET", "/api/users/:id", "Get user profile by ID", "No"),
]
for i, (m, ep, desc, auth) in enumerate(routes_data):
    t.rows[i+1].cells[0].text = m
    t.rows[i+1].cells[1].text = ep
    t.rows[i+1].cells[2].text = desc
    t.rows[i+1].cells[3].text = auth

doc.add_paragraph()

add_para("Registration Flow:", bold=True)
add_bullet("Validates that name, email, username, and password are provided.")
add_bullet("Checks for duplicate email or username in the database.")
add_bullet("Hashes the password using bcrypt (salt rounds = 8).")
add_bullet("Inserts the new user with role='customer'.")

add_para("Login Flow:", bold=True)
add_bullet("Validates username and password are provided.")
add_bullet("Looks up the user by username in the database.")
add_bullet("Compares the provided password against the stored bcrypt hash.")
add_bullet("Returns a safe user object (id, name, email, username, role) without the password.")
add_bullet("The frontend stores user data in localStorage for session management.")

user_routes = open(os.path.join(os.path.dirname(__file__), "backend", "routes", "userRoutes.js"), "r", encoding="utf-8").read()
add_code_block("backend/routes/userRoutes.js", user_routes, "javascript")

doc.add_page_break()

# 6.2 orderRoutes.js
add_heading_custom("6.2 orderRoutes.js – Order Management", level=2)
add_para("File: backend/routes/orderRoutes.js", bold=True)
add_para("This module handles order creation, retrieval, and status updates.")

add_para("Endpoints:", bold=True)

t2 = doc.add_table(rows=5, cols=4)
t2.style = 'Table Grid'
for i, h in enumerate(t_headers):
    t2.rows[0].cells[i].text = h
    for p in t2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

order_routes_data = [
    ("POST", "/api/orders/create", "Place a new order with items", "Customer"),
    ("GET", "/api/orders/user/:id", "Get all orders for a user", "Customer"),
    ("GET", "/api/orders", "Get all orders (admin view)", "Admin"),
    ("PUT", "/api/orders/update-status/:id", "Update order status & step", "Admin"),
]
for i, (m, ep, desc, auth) in enumerate(order_routes_data):
    t2.rows[i+1].cells[0].text = m
    t2.rows[i+1].cells[1].text = ep
    t2.rows[i+1].cells[2].text = desc
    t2.rows[i+1].cells[3].text = auth

doc.add_paragraph()

add_para("Order Creation Flow:", bold=True)
add_bullet("Receives user_id, payment_mode, pickup_slot, and an items array from the frontend.")
add_bullet("Each item has: type (clothing category), qty (quantity), services (array), rate (price per unit).")
add_bullet("Calculates the total: quantity × rate × number of services for each item.")
add_bullet("Inserts a record into the orders table with the computed total.")
add_bullet("Bulk-inserts each item into the order_items table linked to the order.")
add_bullet("Returns the order ID and total to the frontend.")

add_para("Admin Status Update:", bold=True)
add_bullet("Accepts a status string and step number.")
add_bullet("Status values: placed → pickup_scheduled → picked → in_process → ready → out_for_delivery → delivered.")
add_bullet("The step number (0-6) corresponds to the visual progress tracker on the frontend.")

order_routes_code = open(os.path.join(os.path.dirname(__file__), "backend", "routes", "orderRoutes.js"), "r", encoding="utf-8").read()
add_code_block("backend/routes/orderRoutes.js", order_routes_code, "javascript")

doc.add_page_break()

# 6.3 feedbackRoutes.js
add_heading_custom("6.3 feedbackRoutes.js – Feedback System", level=2)
add_para("File: backend/routes/feedbackRoutes.js", bold=True)
add_para("This module handles customer feedback submission and retrieval.")

add_para("Endpoints:", bold=True)

t3 = doc.add_table(rows=3, cols=4)
t3.style = 'Table Grid'
for i, h in enumerate(t_headers):
    t3.rows[0].cells[i].text = h
    for p in t3.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

fb_routes = [
    ("POST", "/api/feedback/create", "Submit feedback with rating", "Customer"),
    ("GET", "/api/feedback/all", "Get all feedback (with user names)", "Admin"),
]
for i, (m, ep, desc, auth) in enumerate(fb_routes):
    t3.rows[i+1].cells[0].text = m
    t3.rows[i+1].cells[1].text = ep
    t3.rows[i+1].cells[2].text = desc
    t3.rows[i+1].cells[3].text = auth

doc.add_paragraph()

add_para("Feedback Flow:", bold=True)
add_bullet("Customer submits: user_id, message (required), rating (1-5, optional), order_id (optional).")
add_bullet("Admin view fetches all feedback with a JOIN to get the user's name.")

fb_code = open(os.path.join(os.path.dirname(__file__), "backend", "routes", "feedbackRoutes.js"), "r", encoding="utf-8").read()
add_code_block("backend/routes/feedbackRoutes.js", fb_code, "javascript")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. FRONTEND PAGES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("7. Frontend Pages", level=1)

add_para(
    "The frontend consists of 10 HTML files served as static pages by the Express server. "
    "Each page includes inline CSS styling and inline JavaScript for interactivity. "
    "The design uses a consistent cream/brown color palette with a photographic background image, "
    "glassmorphism effects (semi-transparent backgrounds with backdrop-blur), and smooth CSS animations."
)

# 7.1 index.html
add_heading_custom("7.1 index.html – Entry Redirect", level=2)
add_para(
    "This is a minimal HTML page that serves as the application entry point. "
    "It uses a <meta http-equiv='refresh'> tag to immediately redirect the user to login.html. "
    "This ensures that when a user visits the root URL (http://localhost:5000/), they are "
    "automatically taken to the login page."
)
add_code_block("frontend/index.html", """<!-- index.html -->
<!DOCTYPE html>
<html>
<head>
    <meta http-equiv="refresh" content="0; url=login.html">
</head>
<body>
</body>
</html>""")

doc.add_paragraph()

# 7.2 login.html
add_heading_custom("7.2 login.html – User Login Page", level=2)
add_para("This page provides a login form for existing users.")

add_para("UI Components:", bold=True)
add_bullet("A centered login box with semi-transparent white background (glassmorphism).")
add_bullet("Username and password input fields with focus effects.")
add_bullet("Login button with hover animation (scale + color change).")
add_bullet("Link to register.html for new users.")

add_para("JavaScript Logic:", bold=True)
add_bullet("loginUser() function sends a POST request to /api/users/login with username and password.")
add_bullet("On success, stores userId, role, and username in localStorage.")
add_bullet("Redirects admin users to admin.html and customers to home.html.")
add_bullet("Displays an alert on invalid credentials.")

add_para("Session Management:", bold=True)
add_bullet("Uses localStorage (not cookies or JWT) for session persistence.")
add_bullet("Stored keys: userId, role, username.")

doc.add_paragraph()

# 7.3 register.html
add_heading_custom("7.3 register.html – User Registration Page", level=2)
add_para("This page provides a registration form for new users.")

add_para("Form Fields:", bold=True)
add_bullet("Username, Email, Full Name, Password, Phone, Address.")

add_para("JavaScript Logic:", bold=True)
add_bullet("registerUser() function sends a POST request to /api/users/register.")
add_bullet("On success, shows a confirmation alert and redirects to login.html.")
add_bullet("On failure, displays the error message from the server.")

doc.add_paragraph()

# 7.4 home.html
add_heading_custom("7.4 home.html – Home / Landing Page", level=2)
add_para("This is the main landing page shown after a customer logs in.")

add_para("UI Components:", bold=True)
add_bullet("Sticky navbar with FreshFold logo, brand name, and navigation links (Home, Services, Order Status, Feedback, Logout).")
add_bullet("Hero section with a personalized welcome message and 'Book Now' button.")
add_bullet("Service cards showcasing Washing, Ironing, and Dry Cleaning with icons from Flaticon CDN.")
add_bullet("Footer with company info, quick links, contact details, and social media icons.")

add_para("JavaScript Logic:", bold=True)
add_bullet("Reads the loggedInUser key from localStorage to display a personalized welcome message.")

doc.add_paragraph()

# 7.5 services.html
add_heading_custom("7.5 services.html – Service Selection & Cart", level=2)
add_para("This is the main ordering page where customers select clothing types, quantities, and services.")

add_para("Clothing Categories (Accordion Cards):", bold=True)
add_bullet("Cotton Wear (👕) – Rate: ₹40/unit")
add_bullet("Silk & Delicate Wear (🥻) – Rate: ₹60/unit")
add_bullet("Denim & Heavy Garments (👖) – Rate: ₹50/unit")
add_bullet("Woolens & Winter Wear (🧥) – Rate: ₹70/unit")

add_para("Each card expands to show:", bold=True)
add_bullet("Quantity controls (+/- buttons with a read-only input field).")
add_bullet("Service checkboxes: Washing, Ironing, Dry Cleaning.")

add_para("JavaScript Functions:", bold=True)
add_bullet("toggleCard(id, event) – Accordion toggle; collapses other cards when one opens.")
add_bullet("changeQty(id, delta) – Increments/decrements quantity, minimum 0.")
add_bullet("updateSummary() – Dynamically rebuilds the order summary section showing selected items.")
add_bullet("generateReceipt() – Validates at least one item is selected, stores the order data in localStorage as 'freshFoldOrder', and redirects to receipt.html.")

doc.add_paragraph()

# 7.6 receipt.html
add_heading_custom("7.6 receipt.html – Order Receipt & Checkout", level=2)
add_para("This page displays an itemized bill before the order is placed.")

add_para("UI Components:", bold=True)
add_bullet("Monospaced font (Courier New) for a receipt-like appearance.")
add_bullet("Table showing: Item Type, Quantity, Services, Amount per item.")
add_bullet("Totals section: Subtotal, GST (5%), Delivery (₹30), Grand Total.")
add_bullet("'Place Order' button.")

add_para("Billing Calculation:", bold=True)
add_bullet("Amount per item = Quantity × Rate × Number of Services Selected")
add_bullet("GST = 5% of subtotal")
add_bullet("Delivery charge = Fixed ₹30")
add_bullet("Grand Total = Subtotal + GST + Delivery")

add_para("JavaScript Logic:", bold=True)
add_bullet("Reads the order from localStorage ('freshFoldOrder').")
add_bullet("Renders the billing table dynamically.")
add_bullet("placeOrder() sends a POST to /api/orders/create with user_id, items array, payment_mode (COD), and pickup_slot.")
add_bullet("On success, stores the orderId in localStorage and redirects to order-status.html.")

doc.add_paragraph()

# 7.7 payment.html
add_heading_custom("7.7 payment.html – Payment Processing Page", level=2)
add_para("This page allows users to select and complete a payment method.")

add_para("Payment Methods:", bold=True)
add_bullet("UPI (Google Pay, PhonePe, Paytm) – Requires UPI ID input.")
add_bullet("Credit/Debit Card – Requires Card Number, Expiry Date, CVV.")
add_bullet("Cash on Delivery (COD) – No additional input needed.")

add_para("JavaScript Logic:", bold=True)
add_bullet("showDetails() dynamically renders input fields based on the selected payment method.")
add_bullet("processPayment() validates the entered payment details and simulates a successful payment.")
add_bullet("Calculates and displays the total amount (Subtotal + 5% GST + ₹30 delivery).")

doc.add_paragraph()

# 7.8 order-status.html
add_heading_custom("7.8 order-status.html – Order Tracking Page", level=2)
add_para("This page shows customers the real-time status of their latest order.")

add_para("UI Components:", bold=True)
add_bullet("Order info card showing: Order ID, current Status, and Last Updated timestamp.")
add_bullet("Visual step-by-step progress tracker with circular indicators.")
add_bullet("Active steps are highlighted with the brand color and a scale animation.")
add_bullet("'Back to Home' button.")

add_para("Tracking Steps (0-6):", bold=True)
add_bullet("0: Order Placed")
add_bullet("1: Pickup Scheduled")
add_bullet("2: Picked Up")
add_bullet("3: In Process")
add_bullet("4: Ready")
add_bullet("5: Out for Delivery")
add_bullet("6: Delivered")

add_para("JavaScript Logic:", bold=True)
add_bullet("loadOrderStatus() fetches the user's orders from /api/orders/user/:id.")
add_bullet("Displays the most recent order's status.")
add_bullet("updateStatusUI(step) highlights the correct steps in the progress tracker.")

doc.add_paragraph()

# 7.9 feedback.html
add_heading_custom("7.9 feedback.html – Customer Feedback Page", level=2)
add_para("This page allows logged-in customers to submit feedback with a star rating.")

add_para("UI Components:", bold=True)
add_bullet("Feedback card with name and email inputs.")
add_bullet("Interactive 5-star rating system (click to select stars).")
add_bullet("Textarea for the feedback message.")
add_bullet("Submit button with hover animation.")

add_para("JavaScript Logic:", bold=True)
add_bullet("Star click handler: toggles active class on stars up to the clicked value.")
add_bullet("Form submit handler: validates that user is logged in (checks localStorage), rating is selected, and message is entered.")
add_bullet("Sends a POST request to /api/feedback/create with user_id, message, and rating.")
add_bullet("Resets the form and stars after successful submission.")

doc.add_paragraph()

# 7.10 admin.html
add_heading_custom("7.10 admin.html – Admin Dashboard", level=2)
add_para("This page provides an admin interface for managing orders and viewing feedback.")

add_para("Tab Navigation:", bold=True)
add_bullet("Orders Tab – Displays all customer orders in a table.")
add_bullet("Feedbacks Tab – Displays all customer feedback.")

add_para("Orders Table Columns:", bold=True)
add_bullet("Order ID, Customer Name, Total Amount (₹), Status (dropdown), Payment Mode, Pickup Slot, Update (Save button).")

add_para("Order Status Management:", bold=True)
add_bullet("Each order row has a dropdown with all possible statuses: placed, pickup_scheduled, picked, in_process, ready, out_for_delivery, delivered, cancelled.")
add_bullet("The admin selects a new status and clicks 'Save' to update.")
add_bullet("updateStatus() sends a PUT request to /api/orders/update-status/:id with the new status and corresponding step number.")

add_para("Feedbacks Table Columns:", bold=True)
add_bullet("Feedback ID, Customer Name, Order ID, Message, Rating.")

add_para("Logout:", bold=True)
add_bullet("Clears userId, role, and username from localStorage and redirects to login.html.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. APPLICATION FLOW
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("8. Application Flow & User Journey", level=1)

add_heading_custom("Customer Journey", level=2)
add_para("1. User visits the application → Redirected to Login page.")
add_para("2. New users click 'Sign up' → Fill registration form → Account created → Redirected to Login.")
add_para("3. User logs in with username & password → Credentials verified → Redirected to Home page.")
add_para("4. User clicks 'Book Now' or 'Services' → Service Selection page.")
add_para("5. User selects clothing types, sets quantities, checks services → Clicks 'Proceed'.")
add_para("6. Receipt page shows itemized bill with GST and delivery → User clicks 'Place Order'.")
add_para("7. Order is saved to the database → User is redirected to Order Status page.")
add_para("8. User can track their order progress in real-time.")
add_para("9. User can submit feedback with a star rating from the Feedback page.")
add_para("10. User clicks 'Logout' → localStorage cleared → Redirected to Login.")

add_heading_custom("Admin Journey", level=2)
add_para("1. Admin logs in with admin credentials → Redirected to Admin Dashboard.")
add_para("2. Orders tab shows all customer orders with current status.")
add_para("3. Admin updates order status via dropdown → Clicks 'Save' → Database updated.")
add_para("4. Admin switches to Feedbacks tab to view all customer feedback and ratings.")
add_para("5. Admin clicks 'Logout' to end their session.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. API ENDPOINTS REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("9. API Endpoints Reference", level=1)

api_table = doc.add_table(rows=8, cols=5)
api_table.style = 'Table Grid'
api_headers = ["#", "Method", "Endpoint", "Description", "Request Body"]
for i, h in enumerate(api_headers):
    api_table.rows[0].cells[i].text = h
    for p in api_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

api_data = [
    ("1", "POST", "/api/users/register", "Register new user", "{name, email, username, password, phone, address}"),
    ("2", "POST", "/api/users/login", "Login user", "{username, password}"),
    ("3", "GET", "/api/users/:id", "Get user by ID", "—"),
    ("4", "POST", "/api/orders/create", "Create new order", "{user_id, items[], payment_mode, pickup_slot}"),
    ("5", "GET", "/api/orders/user/:id", "Get user's orders", "—"),
    ("6", "GET", "/api/orders", "Get all orders (admin)", "—"),
    ("7", "PUT", "/api/orders/update-status/:id", "Update order status", "{status, step}"),
]
for i, row_data in enumerate(api_data):
    for j, val in enumerate(row_data):
        api_table.rows[i+1].cells[j].text = val

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. ENVIRONMENT & CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("10. Environment & Configuration", level=1)

add_para(
    "The backend uses a .env file (in the backend/ directory) for sensitive configuration. "
    "This file is excluded from version control via .gitignore."
)

add_para("Expected .env Variables:", bold=True)
add_code_block(".env", """DB_HOST=localhost
DB_USER=root
DB_PASSWORD=your_mysql_password
DB_NAME=freshfold
PORT=5000""")

add_para(".gitignore Rules:", bold=True)
add_bullet("node_modules/ – Excluded to keep the repository lightweight.")
add_bullet(".env – Excluded to protect sensitive credentials.")
add_bullet("dist/, build/ – Excluded build output.")
add_bullet(".DS_Store, .vscode/, .idea/ – OS and IDE files excluded.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 11. HOW TO RUN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("11. How to Run the Project", level=1)

add_heading_custom("Prerequisites", level=2)
add_bullet("Node.js (v14 or later) installed")
add_bullet("MySQL Server running locally")
add_bullet("A MySQL client (MySQL Workbench, phpMyAdmin, or CLI)")

add_heading_custom("Step 1: Set Up the Database", level=2)
add_para("Open your MySQL client and run the db.sql file to create the database, tables, and seed the services catalog.")
add_code_block("MySQL Command", "mysql -u root -p < db.sql")

add_heading_custom("Step 2: Configure Environment Variables", level=2)
add_para("Create a .env file inside the backend/ directory with your MySQL credentials.")

add_heading_custom("Step 3: Install Dependencies", level=2)
add_code_block("Terminal", """cd backend
npm install""")

add_heading_custom("Step 4: Seed the Admin User", level=2)
add_code_block("Terminal", "npm run seed")
add_para("This creates an admin user: username='admin1', password='12345'.")

add_heading_custom("Step 5: Start the Server", level=2)
add_code_block("Terminal", "npm start")
add_para("The server starts on http://localhost:5000. Open this URL in your browser.")

add_heading_custom("Step 6: Access the Application", level=2)
add_bullet("Customer Login: Use credentials created via the registration page.")
add_bullet("Admin Login: Use username='admin1' and password='12345'.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 12. KEY FEATURES SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

add_heading_custom("12. Key Features Summary", level=1)

features_table = doc.add_table(rows=13, cols=2)
features_table.style = 'Table Grid'
features_table.rows[0].cells[0].text = "Feature"
features_table.rows[0].cells[1].text = "Details"
for p in features_table.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in features_table.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True

features = [
    ("User Authentication", "Register/Login with bcrypt password hashing"),
    ("Role-based Access", "Customer and Admin roles with separate dashboards"),
    ("Service Catalog", "4 clothing categories with 3 service options each"),
    ("Interactive Cart", "Accordion cards, quantity controls, service checkboxes"),
    ("Order Receipt", "Itemized billing with GST and delivery charges"),
    ("Multiple Payments", "UPI, Card, and Cash on Delivery options"),
    ("Order Tracking", "7-step visual progress tracker"),
    ("Feedback System", "5-star rating with message submission"),
    ("Admin Dashboard", "Order management and feedback viewing"),
    ("Responsive Design", "Mobile-friendly with media queries"),
    ("Glassmorphism UI", "Semi-transparent backgrounds with blur effects"),
    ("REST API", "Clean separation of frontend and backend via API"),
]

for i, (feat, detail) in enumerate(features):
    features_table.rows[i+1].cells[0].text = feat
    features_table.rows[i+1].cells[1].text = detail

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Documentation —")
set_run_font(run, size=14, bold=True, color=(90, 58, 50))


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), "FreshFold_Project_Documentation.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
